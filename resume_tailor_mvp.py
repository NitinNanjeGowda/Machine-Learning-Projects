# -----------------------------
# Resume Tailor MVP (Local)
# -----------------------------

from docx import Document
import spacy
import re
from collections import Counter
from sentence_transformers import SentenceTransformer, util
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline

# -----------------------------
# Step 0: Load .docx
# -----------------------------
def load_docx(path):
    doc = Document(path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
    return text

# -----------------------------
# Step 1: Clean text
# -----------------------------
def clean_text(text):
    text = text.lower()
    text = re.sub(r"[^a-z0-9\s]", " ", text)
    return text

# -----------------------------
# Step 2: Extract keywords from job description
# -----------------------------
nlp = spacy.load("en_core_web_sm")

def extract_keywords(text, top_n=50):
    text = clean_text(text)
    doc = nlp(text)
    keywords = [token.text for token in doc if token.pos_ in ["NOUN", "PROPN"] and len(token.text)>2]
    counts = Counter(keywords)
    return [word for word, freq in counts.most_common(top_n)]

# -----------------------------
# Step 3: Rank resume sentences using embeddings
# -----------------------------
embed_model = SentenceTransformer('all-MiniLM-L6-v2')  # CPU-friendly

def rank_resume_sentences(resume_text, job_text):
    resume_sentences = [s.strip() for s in resume_text.split("\n") if s.strip()]
    resume_embs = embed_model.encode(resume_sentences)
    job_emb = embed_model.encode(job_text)
    similarities = [util.cos_sim(job_emb, emb).item() for emb in resume_embs]
    ranked = sorted(zip(similarities, resume_sentences), key=lambda x: x[0], reverse=True)
    return [s for sim, s in ranked]

# -----------------------------
# Step 4: Rewrite sentences to include missing keywords
# -----------------------------
# CPU-friendly local model
model_name = "google/flan-t5-base"  
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
generator = pipeline("text2text-generation", model=model, tokenizer=tokenizer)

def rewrite_sentence(sentence, keywords):
    prompt = f"Rewrite this resume sentence to naturally include these keywords for ATS: {', '.join(keywords)}.\nSentence: {sentence}"
    output = generator(prompt, max_length=100)[0]['generated_text']
    return output

# -----------------------------
# Step 5: Assemble tailored resume
# -----------------------------
def tailor_resume(resume_text, job_text):
    jd_keywords = extract_keywords(job_text)
    ranked_sentences = rank_resume_sentences(resume_text, job_text)

    tailored_sentences = []
    for sent in ranked_sentences:
        sent_clean = clean_text(sent).split()
        missing = [kw for kw in jd_keywords if kw not in sent_clean]
        if missing:
            sent = rewrite_sentence(sent, missing[:5])  # include up to 5 missing keywords per sentence
        tailored_sentences.append(sent)

    # Add missing skills section
    tailored_sentences.append("\nSkills for ATS: " + ", ".join(jd_keywords[:20]))
    return tailored_sentences

# -----------------------------
# Step 6: Save as .docx
# -----------------------------
def save_resume(sentences, filename="tailored_resume.docx"):
    doc = Document()
    doc.add_heading("Tailored Resume", 0)
    for s in sentences:
        doc.add_paragraph(s)
    doc.save(filename)

def ats_score_weighted(resume_text, job_text, top_n=50):
    """
    Compute weighted ATS score based on top N keywords from JD.
    Score = % of keywords present in resume
    """
    jd_keywords = extract_keywords(job_text, top_n)
    resume_words = set(clean_text(resume_text).split())
    
    matched = [kw for kw in jd_keywords if kw in resume_words]
    score = (len(matched) / len(jd_keywords)) * 100
    return round(score, 2), matched, [kw for kw in jd_keywords if kw not in resume_words]

def iterative_tailor(resume_text, job_text, target_score=90, max_iter=5):
    current_resume = resume_text
    for i in range(max_iter):
        score, matched, missing = ats_score_weighted(current_resume, job_text)
        print(f"Iteration {i+1}: ATS Score = {score}% | Missing Keywords: {len(missing)}")
        
        if score >= target_score or not missing:
            break
        
        # Rewrite top sentences to include missing keywords
        sentences = [s.strip() for s in current_resume.split("\n") if s.strip()]
        for idx, sent in enumerate(sentences):
            sent_clean = clean_text(sent).split()
            # Only rewrite if it can include missing keywords
            kw_to_add = [kw for kw in missing if kw not in sent_clean]
            if kw_to_add:
                sentences[idx] = rewrite_sentence(sent, kw_to_add[:5])  # rewrite with up to 5 missing keywords
        
        # Reassemble resume
        current_resume = "\n".join(sentences)
    
    # Final tailored resume with a skills section
    final_sentences = current_resume.split("\n")
    final_sentences.append("\nSkills for ATS: " + ", ".join(extract_keywords(job_text, 20)))
    
    return final_sentences, score

# -----------------------------
# MAIN
# -----------------------------
if __name__ == "__main__":
    resume_path = "Nitin Nanje Gowda Resume.docx"               # your resume file
    job_path = "Job Description.docx"        # your job description file

    print("Loading files...")
    resume_text = load_docx(resume_path)
    job_text = load_docx(job_path)

    print("Iteratively tailoring resume for ATS...")
    tailored_sentences, final_score = iterative_tailor(resume_text, job_text, target_score=90)

    save_resume(tailored_sentences)
    print(f"✅ Tailored resume saved as 'tailored_resume.docx' | Final ATS Score: {final_score}%")
